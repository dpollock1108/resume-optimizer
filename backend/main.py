import json
import io
import logging
import os
from pathlib import Path
from typing import Optional
from uuid import uuid4

from fastapi import Depends, FastAPI, HTTPException, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, JSONResponse, StreamingResponse
from fastapi.staticfiles import StaticFiles
from google.api_core.exceptions import GoogleAPICallError
from pydantic import BaseModel, Field
from dotenv import load_dotenv
from anthropic import Anthropic
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from auth import AuthenticatedUser, get_authenticated_user
from database import (
    create_document,
    delete_document,
    ensure_user,
    list_documents,
    update_document,
)
from rate_limit import enforce_ai_rate_limit

load_dotenv()
load_dotenv(".env.local")
logger = logging.getLogger(__name__)

app = FastAPI(title="Resume Optimizer")
ANTHROPIC_MODEL = os.getenv("ANTHROPIC_MODEL", "claude-sonnet-5")


@app.exception_handler(GoogleAPICallError)
async def google_api_error(request: Request, error: GoogleAPICallError):
    logger.error(
        "Google Cloud request failed for %s",
        request.url.path,
        exc_info=(type(error), error, error.__traceback__),
    )
    return JSONResponse(
        status_code=503,
        content={
            "detail": "Database unavailable. Check the local Firestore project and API configuration."
        },
    )

if os.getenv("APP_ENV", "development") == "development":
    app.add_middleware(
        CORSMiddleware,
        allow_origins=["http://localhost:5173"],
        allow_credentials=True,
        allow_methods=["*"],
        allow_headers=["*"],
    )

def anthropic_client() -> Anthropic:
    api_key = os.getenv("ANTHROPIC_API_KEY", "").strip()
    return Anthropic(api_key=api_key or None)


def extract_text(message) -> str:
    """Return the first text block's content.

    Models with extended thinking (e.g. claude-sonnet-5) emit a ThinkingBlock
    as content[0], so we can't assume index 0 is text.
    """
    for block in message.content:
        if getattr(block, "type", None) == "text":
            return block.text
    raise ValueError("No text block in Anthropic response")


def parse_json_response(message) -> dict:
    """Parse a JSON object out of an Anthropic response.

    Tolerates markdown ``` / ```json fences and any leading/trailing prose by
    falling back to the outermost {...} span. Logs the raw text on failure so
    truncated or malformed responses are debuggable.
    """
    raw = extract_text(message).strip()

    if raw.startswith("```"):
        raw = raw.split("\n", 1)[-1]
        if raw.rstrip().endswith("```"):
            raw = raw.rstrip()[:-3]
    raw = raw.strip()

    # strict=False allows literal newlines/tabs inside strings — the model emits
    # real line breaks inside the tailored_resume value rather than escaping them.
    try:
        return json.loads(raw, strict=False)
    except json.JSONDecodeError:
        start, end = raw.find("{"), raw.rfind("}")
        if start != -1 and end > start:
            try:
                return json.loads(raw[start : end + 1], strict=False)
            except json.JSONDecodeError:
                pass
        logger.error(
            "Failed to parse Anthropic JSON (stop_reason=%s, len=%d): %.500s",
            getattr(message, "stop_reason", "?"),
            len(raw),
            raw,
        )
        raise HTTPException(status_code=502, detail="Failed to parse API response as JSON")


class AnalyzeRequest(BaseModel):
    resume: str = Field(min_length=1, max_length=50_000)
    job_posting: str = Field(min_length=1, max_length=50_000)
    notes: Optional[str] = Field(default=None, max_length=5_000)
    output_format: str = "classic"
    template_hints: Optional[str] = Field(default=None, max_length=10_000)


class BulletEdit(BaseModel):
    original: str
    revised: str
    reason: str


class AnalyzeResponse(BaseModel):
    before_match_score: int
    before_ats_score: int
    match_score: int
    ats_score: int
    effort_score: int
    matched_keywords: list[str]
    missing_keywords: list[str]
    tailored_summary: str
    bullet_edits: list[BulletEdit]
    cover_opener: str
    tailored_resume: str
    strategic_note: str


class ExportRequest(BaseModel):
    tailored_resume: str = Field(min_length=1, max_length=75_000)
    layout: str = "classic"
    name: str = Field(default="", max_length=200)


class UserResponse(BaseModel):
    id: str
    email: str


class ProfileCreate(BaseModel):
    name: str = Field(min_length=1, max_length=200)
    resume: str = Field(min_length=1, max_length=50_000)


class ProfileResponse(ProfileCreate):
    id: str


@app.get("/healthz", include_in_schema=False)
async def healthz():
    return {"status": "ok"}


@app.get("/api/me", response_model=UserResponse)
async def current_user(user: AuthenticatedUser = Depends(get_authenticated_user)):
    ensure_user(user)
    return UserResponse(id=user.id, email=user.email)


@app.get("/api/profiles", response_model=list[ProfileResponse])
async def get_profiles(user: AuthenticatedUser = Depends(get_authenticated_user)):
    ensure_user(user)
    return list_documents(user, "profiles")


@app.post("/api/profiles", response_model=ProfileResponse, status_code=201)
async def create_profile(
    profile: ProfileCreate,
    user: AuthenticatedUser = Depends(get_authenticated_user),
):
    ensure_user(user)
    return create_document(user, "profiles", str(uuid4()), profile.model_dump())


ANALYZE_PROMPT = """You are an expert resume consultant and ATS optimization specialist.

Given a candidate's resume and a job posting, analyze the fit and produce a fully tailored, ready-to-use version.

<resume>
{resume}
</resume>

<job_posting>
{job_posting}
</job_posting>

{notes_section}

{template_section}

Output format requested: {output_format}

IMPORTANT INSTRUCTIONS FOR THE TAILORED RESUME:
- The "tailored_resume" field must be the COMPLETE, FINAL resume with ALL edits already applied — not the original with suggestions on the side.
- Rewrite the professional summary to target this specific role.
- Rephrase bullet points to incorporate keywords from the job posting where the candidate's experience genuinely supports them. Massage the language to align with the job's terminology and priorities.
- Reorder skills to lead with the most relevant ones for this role.
- Do NOT fabricate experience, skills, or accomplishments the candidate doesn't have. Only reshape and reframe what's already there.
- The user should be able to use the tailored resume as-is with zero additional editing.
- Preserve the candidate's name, contact info, dates, company names, and job titles exactly as they appear in the original.

Return a JSON object with exactly these fields:
- "before_match_score": integer 0-100 — keyword/skill match % of the ORIGINAL resume against this job posting (before any edits)
- "before_ats_score": integer 0-100 — ATS pass likelihood of the ORIGINAL resume (before any edits)
- "match_score": integer 0-100 — keyword/skill match % of the TAILORED resume (after your edits)
- "ats_score": integer 0-100 — ATS pass likelihood of the TAILORED resume (after your edits)
- "effort_score": integer 0-100 representing how much of the base resume is directly reusable (high = less work needed)
- "matched_keywords": array of up to 12 strings — keywords/skills found in both resume and job posting
- "missing_keywords": array of up to 8 strings — important keywords from the job posting not found in the resume
- "tailored_summary": string — the 2-3 sentence rewritten professional summary you used in the tailored resume
- "bullet_edits": array of objects, each with "original" (the exact original bullet from the resume), "revised" (your improved version as it appears in the tailored resume), and "reason" (brief explanation of the change). Include 3-8 of the most impactful edits. These should reflect changes you ALREADY MADE in the tailored_resume.
- "cover_opener": string — a compelling 2-3 sentence cover letter opening paragraph tailored to this role
- "tailored_resume": string — the COMPLETE tailored resume as clean plain text with all edits applied. Formatting rules: put the candidate's name on line 1, all contact info on a SINGLE line 2 (no line breaks within contact info), then sections with UPPERCASE headers. Use "•" for bullet points. Use exactly ONE blank line between sections — never more.
- "strategic_note": string — one paragraph of honest, direct strategic advice about the candidate's fit for this role and how to position themselves

Be specific, actionable, and honest. Don't inflate scores — if the fit is poor, say so constructively.

Return ONLY valid JSON, no markdown fences or other text."""


@app.post("/api/analyze", response_model=AnalyzeResponse)
async def analyze(
    req: AnalyzeRequest,
    user: AuthenticatedUser = Depends(get_authenticated_user),
):
    enforce_ai_rate_limit(user.id)
    notes_section = f"<additional_notes>\n{req.notes}\n</additional_notes>" if req.notes else ""
    template_section = (
        f"<template_conventions>\nMatch this template structure:\n{req.template_hints}\n</template_conventions>"
        if req.template_hints
        else ""
    )

    prompt = ANALYZE_PROMPT.format(
        resume=req.resume,
        job_posting=req.job_posting,
        notes_section=notes_section,
        template_section=template_section,
        output_format=req.output_format,
    )

    try:
        message = anthropic_client().messages.create(
            model=ANTHROPIC_MODEL,
            max_tokens=8192,
            messages=[{"role": "user", "content": prompt}],
        )
    except Exception:
        logger.exception("Anthropic analysis request failed for user %s", user.id)
        raise HTTPException(status_code=502, detail="AI analysis request failed") from None

    data = parse_json_response(message)

    return AnalyzeResponse(**data)


# --- Interview Prep ---


class PrepRequest(BaseModel):
    resume: str = Field(min_length=1, max_length=50_000)
    job_posting: str = Field(min_length=1, max_length=50_000)
    company: str = Field(default="", max_length=500)
    role: str = Field(default="", max_length=500)


class KnowledgeGap101(BaseModel):
    topic: str
    why_it_matters: str
    key_concepts: list[str]
    quick_study: str


class PrepResponse(BaseModel):
    company_overview: str
    role_breakdown: str
    likely_questions: list[str]
    talking_points: list[str]
    knowledge_gaps: list[KnowledgeGap101]
    culture_notes: str


class ApplicationCreate(BaseModel):
    id: str = Field(min_length=1, max_length=100)
    profile_id: str = Field(min_length=1, max_length=100)
    profile_name: str = Field(min_length=1, max_length=200)
    company: str = Field(min_length=1, max_length=500)
    role: str = Field(min_length=1, max_length=500)
    job_posting: str = Field(min_length=1, max_length=50_000)
    status: str = Field(pattern="^(applied|screening|interviewing|offer|rejected|withdrawn)$")
    date_applied: str = Field(max_length=100)
    match_score: int = Field(ge=0, le=100)
    ats_score: int = Field(ge=0, le=100)
    tailored_resume: str = Field(min_length=1, max_length=75_000)
    layout: str = Field(pattern="^(classic|modern|compact|custom)$")
    notes: str = Field(default="", max_length=10_000)
    prep: Optional[PrepResponse] = None


class ApplicationUpdate(BaseModel):
    status: Optional[str] = Field(
        default=None,
        pattern="^(applied|screening|interviewing|offer|rejected|withdrawn)$",
    )
    notes: Optional[str] = Field(default=None, max_length=10_000)
    prep: Optional[PrepResponse] = None


class ApplicationResponse(ApplicationCreate):
    pass


@app.get("/api/applications", response_model=list[ApplicationResponse])
async def get_applications(user: AuthenticatedUser = Depends(get_authenticated_user)):
    return list_documents(user, "applications")


@app.post("/api/applications", response_model=ApplicationResponse, status_code=201)
async def create_application(
    application: ApplicationCreate,
    user: AuthenticatedUser = Depends(get_authenticated_user),
):
    payload = application.model_dump()
    document_id = payload.pop("id")
    return create_document(user, "applications", document_id, payload)


@app.patch("/api/applications/{application_id}", response_model=ApplicationResponse)
async def patch_application(
    application_id: str,
    updates: ApplicationUpdate,
    user: AuthenticatedUser = Depends(get_authenticated_user),
):
    updated = update_document(
        user,
        "applications",
        application_id,
        updates.model_dump(exclude_unset=True),
    )
    if updated is None:
        raise HTTPException(status_code=404, detail="Application not found")
    return updated


@app.delete("/api/applications/{application_id}", status_code=204)
async def remove_application(
    application_id: str,
    user: AuthenticatedUser = Depends(get_authenticated_user),
):
    if not delete_document(user, "applications", application_id):
        raise HTTPException(status_code=404, detail="Application not found")


PREP_PROMPT = """You are an expert interview coach and career strategist.

Given a candidate's resume and a job posting, generate a comprehensive interview prep guide.

<resume>
{resume}
</resume>

<job_posting>
{job_posting}
</job_posting>

<company>{company}</company>
<role>{role}</role>

Your job is to help this candidate prepare for an interview for this specific role. Be practical, specific, and honest about gaps.

Return a JSON object with exactly these fields:

- "company_overview": string — 2-3 sentences about the company, what they do, their industry position, and anything notable a candidate should know walking into an interview.

- "role_breakdown": string — 2-3 sentences explaining what this role actually involves day-to-day, what the hiring manager likely cares about most, and what success looks like in the first 6 months.

- "likely_questions": array of 8-10 strings — specific interview questions the candidate is likely to face. Mix behavioral ("Tell me about a time..."), technical, and role-specific questions. Tailor these to the actual job posting, not generic questions.

- "talking_points": array of 5-7 strings — specific experiences, projects, or skills from the candidate's resume that they should emphasize and have stories ready for. Each should be a concrete suggestion like "Your work on X at Y directly maps to their need for Z — prepare a detailed walkthrough."

- "knowledge_gaps": array of objects for areas where the candidate's resume suggests they may lack depth relative to what the job requires. Each object has:
  - "topic": string — the skill or domain (e.g. "Kubernetes orchestration", "SOC 2 compliance")
  - "why_it_matters": string — 1 sentence on why this matters for the role
  - "key_concepts": array of 4-6 strings — the essential terms and concepts to understand
  - "quick_study": string — a focused 2-3 sentence briefing that gives the candidate enough context to speak intelligently about this topic without claiming expertise. Include what it is, how it's used in practice, and one concrete example.

- "culture_notes": string — 2-3 sentences on what the company culture seems like based on the job posting language, and how the candidate should present themselves (formal vs casual, innovation-focused vs process-focused, etc.)

Be honest about gaps — don't skip them to be nice. The candidate needs to know where they're weak so they can prepare. But frame gaps constructively: "You'll want to brush up on X" not "You don't know X."

Include 3-6 knowledge gap topics. Focus on gaps that are likely to come up in the interview, not obscure edge cases.

Return ONLY valid JSON, no markdown fences or other text."""


@app.post("/api/prep", response_model=PrepResponse)
async def prep(
    req: PrepRequest,
    user: AuthenticatedUser = Depends(get_authenticated_user),
):
    enforce_ai_rate_limit(user.id)
    prompt = PREP_PROMPT.format(
        resume=req.resume,
        job_posting=req.job_posting,
        company=req.company or "Unknown",
        role=req.role or "Unknown",
    )

    try:
        message = anthropic_client().messages.create(
            model=ANTHROPIC_MODEL,
            max_tokens=8192,
            messages=[{"role": "user", "content": prompt}],
        )
    except Exception:
        logger.exception("Anthropic prep request failed for user %s", user.id)
        raise HTTPException(status_code=502, detail="AI prep request failed") from None

    data = parse_json_response(message)

    return PrepResponse(**data)


def _add_bottom_border(paragraph):
    from docx.oxml.ns import qn
    from lxml import etree
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = etree.SubElement(pPr, qn("w:pBdr"))
    bottom = etree.SubElement(pBdr, qn("w:bottom"))
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1A1A2E")


def _is_section_heading(stripped: str) -> bool:
    if stripped.startswith(("•", "-", "*")):
        return False
    if stripped.isupper() and len(stripped) < 60:
        return True
    return False


def _is_name_line(stripped: str, is_first_content: bool) -> bool:
    if not is_first_content:
        return False
    if len(stripped) < 50 and not any(c in stripped for c in "@.|•-*"):
        return True
    return False


def _clean_lines(resume_text: str) -> list[str]:
    raw = [line.strip() for line in resume_text.split("\n")]
    lines: list[str] = []
    for i, stripped in enumerate(raw):
        if not stripped:
            if lines and not lines[-1]:
                continue
            lines.append("")
        else:
            if _is_section_heading(stripped) and lines and lines[-1] != "":
                lines.append("")
            lines.append(stripped)
    return lines


def build_classic(doc: Document, resume_text: str, name: str):
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    section = doc.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    first_content = True
    for stripped in _clean_lines(resume_text):
        if not stripped:
            p = doc.add_paragraph("")
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            continue

        if _is_name_line(stripped, first_content):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(stripped.upper())
            run.bold = True
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
            p.space_after = Pt(2)
            first_content = False
            continue

        first_content = False

        if _is_section_heading(stripped):
            p = doc.add_paragraph()
            run = p.add_run(stripped.upper())
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
            p.space_before = Pt(16)
            p.space_after = Pt(3)
            _add_bottom_border(p)
        elif stripped.startswith(("•", "-", "*")):
            bullet_text = stripped.lstrip("•-* ")
            p = doc.add_paragraph(bullet_text, style="List Bullet")
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
        else:
            p = doc.add_paragraph(stripped)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)


def build_modern(doc: Document, resume_text: str, name: str):
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10.5)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    first_content = True
    for stripped in _clean_lines(resume_text):
        if not stripped:
            p = doc.add_paragraph("")
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            continue

        if _is_name_line(stripped, first_content):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(stripped)
            run.bold = True
            run.font.size = Pt(22)
            run.font.color.rgb = RGBColor(0x2D, 0x5A, 0x88)
            p.space_after = Pt(2)
            first_content = False
            continue

        first_content = False

        if _is_section_heading(stripped):
            p = doc.add_paragraph()
            run = p.add_run(stripped.upper())
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x2D, 0x5A, 0x88)
            p.space_before = Pt(8)
            p.space_after = Pt(3)
            _add_bottom_border(p)
        elif stripped.startswith(("•", "-", "*")):
            bullet_text = stripped.lstrip("•-* ")
            p = doc.add_paragraph(bullet_text, style="List Bullet")
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
        else:
            p = doc.add_paragraph(stripped)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)


def build_compact(doc: Document, resume_text: str, name: str):
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Arial Narrow"
    font.size = Pt(10)
    font.color.rgb = RGBColor(0x22, 0x22, 0x22)

    section = doc.sections[0]
    section.top_margin = Inches(0.4)
    section.bottom_margin = Inches(0.4)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    first_content = True
    for stripped in _clean_lines(resume_text):
        if not stripped:
            continue

        if _is_name_line(stripped, first_content):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(stripped.upper())
            run.bold = True
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(0x11, 0x11, 0x11)
            p.space_after = Pt(1)
            first_content = False
            continue

        first_content = False

        if _is_section_heading(stripped):
            p = doc.add_paragraph()
            run = p.add_run(stripped.upper())
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x11, 0x11, 0x11)
            p.space_before = Pt(6)
            p.space_after = Pt(1)
            _add_bottom_border(p)
        elif stripped.startswith(("•", "-", "*")):
            bullet_text = stripped.lstrip("•-* ")
            p = doc.add_paragraph(bullet_text, style="List Bullet")
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            for run in p.runs:
                run.font.size = Pt(9.5)
        else:
            p = doc.add_paragraph(stripped)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(1)


BUILDERS = {
    "classic": build_classic,
    "modern": build_modern,
    "compact": build_compact,
}


@app.post("/api/export")
async def export_docx(
    req: ExportRequest,
    user: AuthenticatedUser = Depends(get_authenticated_user),
):
    doc = Document()
    builder = BUILDERS.get(req.layout, build_classic)
    builder(doc, req.tailored_resume, req.name)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    filename = f"resume_{req.layout}.docx"
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


STATIC_DIR = Path(__file__).parent / "static"
if STATIC_DIR.exists():
    assets_directory = STATIC_DIR / "assets"
    if assets_directory.exists():
        app.mount("/assets", StaticFiles(directory=assets_directory), name="assets")

    @app.get("/{path:path}", include_in_schema=False)
    async def frontend(path: str, request: Request):
        get_authenticated_user(request)
        requested_file = STATIC_DIR / path
        if path and requested_file.is_file() and STATIC_DIR in requested_file.resolve().parents:
            return FileResponse(requested_file)
        return FileResponse(STATIC_DIR / "index.html")
